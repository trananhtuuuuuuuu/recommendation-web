"""LayoutLMv3 inference and document text extraction."""

from __future__ import annotations

import csv
import io
import json
import logging
import os
from pathlib import Path
import shutil
import subprocess
import tempfile
from threading import Lock
from typing import Any

from .postprocess import EntitySpan, build_canonical, build_profile


LOGGER = logging.getLogger(__name__)
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"}
MODEL_FILE = "config.json"

# --------------------------------------------------------------------------- #
# PaddleOCR -- the SAME OCR family the LayoutLMv3 model was trained with (the    #
# notebook used PaddleOCR line detection + a char-proportional word split). At  #
# inference we replicate that exactly so the word boxes match training; Tesseract#
# stays as a fallback. paddlepaddle 3.x needs oneDNN + the new PIR executor off  #
# on this CPU, so those flags are forced before the model is built.             #
# --------------------------------------------------------------------------- #
_PADDLE: dict[str, Any] = {"engine": None, "tried": False}


def _load_paddle_engine() -> Any | None:
    """Lazy singleton PaddleOCR engine; returns None if it can't be built."""
    if _PADDLE["tried"]:
        return _PADDLE["engine"]
    _PADDLE["tried"] = True
    os.environ.setdefault("FLAGS_use_mkldnn", "0")
    os.environ.setdefault("FLAGS_enable_pir_in_executor", "0")
    try:
        from paddleocr import PaddleOCR

        _PADDLE["engine"] = PaddleOCR(
            lang=os.getenv("AI_PADDLE_LANG", "en"),
            use_doc_orientation_classify=False,
            use_doc_unwarping=False,
            use_textline_orientation=False,
            enable_mkldnn=False,
        )
    except Exception:
        LOGGER.warning("PaddleOCR unavailable; using Tesseract OCR.", exc_info=True)
        _PADDLE["engine"] = None
    return _PADDLE["engine"]


def _quad_to_aabb(quad: Any) -> list[float]:
    """Axis-aligned [x1,y1,x2,y2] from a PaddleOCR 4-point line polygon."""
    xs = [float(point[0]) for point in quad]
    ys = [float(point[1]) for point in quad]
    return [min(xs), min(ys), max(xs), max(ys)]


def _split_line_to_words(text: str, box: list[float]) -> list[tuple[str, list[float]]]:
    """Split a line box into per-word boxes by character ratio (notebook logic)."""
    x1, y1, x2, y2 = box
    words = text.split()
    if not words:
        return []
    total = max(sum(len(word) for word in words) + max(len(words) - 1, 0), 1)
    out: list[tuple[str, list[float]]] = []
    cursor, span = x1, (x2 - x1)
    for word in words:
        word_x2 = cursor + (len(word) / total) * span
        out.append((word, [cursor, y1, word_x2, y2]))
        cursor = word_x2 + (1.0 / total) * span  # leave one space-width gap
    return out


class UnsupportedDocumentError(ValueError):
    """Raised when the uploaded CV format cannot be parsed."""


class CvParser:
    """Parse CV documents with a trained LayoutLMv3 model and safe fallbacks."""

    def __init__(
        self,
        model_dir: str,
        max_pages: int = 3,
        enable_model: bool = False,
        ocr_languages: str = "eng+vie",
        ocr_engine: str = "tesseract",
    ) -> None:
        self.model_dir = Path(model_dir)
        self.max_pages = max(1, max_pages)
        self.enable_model = enable_model
        self.ocr_languages = ocr_languages
        # "paddle" matches the training OCR (falls back to Tesseract if it fails);
        # "tesseract" is the lightweight default.
        self.ocr_engine = ocr_engine
        self._model_lock = Lock()
        self._processor: Any = None
        self._model: Any = None
        self._torch: Any = None

    @property
    def model_available(self) -> bool:
        """Return whether model dependencies and notebook output are available."""
        return self.enable_model and (self.model_dir / MODEL_FILE).is_file()

    @property
    def image_ocr_available(self) -> bool:
        """Return whether lightweight image OCR is installed."""
        return shutil.which("tesseract") is not None

    def parse(self, content: bytes, filename: str) -> dict:
        """Parse an uploaded CV and return legacy frontend-ready profile fields."""
        entities, document_text, model_used = self._collect(content, filename)
        return build_profile(entities, document_text, model_used)

    def parse_canonical(self, content: bytes, filename: str) -> dict:
        """Parse an uploaded CV into the canonical structure (recommender input)."""
        entities, document_text, model_used = self._collect(content, filename)
        return build_canonical(entities, document_text, model_used)

    def _collect(self, content: bytes, filename: str) -> tuple[list[EntitySpan], str, bool]:
        """Extract entity spans + document text once, shared by both builders."""
        suffix = Path(filename or "").suffix.lower()
        if suffix not in IMAGE_EXTENSIONS | {".pdf", ".doc", ".docx"}:
            raise UnsupportedDocumentError(
                "CV must be a PDF, DOC, DOCX, PNG, JPG, WebP, BMP, or TIFF file."
            )

        document_text = self._extract_text(content, suffix)
        entities: list[EntitySpan] = []
        model_used = False

        if self.model_available and suffix in IMAGE_EXTENSIONS | {".pdf"}:
            pages = self._load_visual_pages(content, suffix)
            for page_index, page in enumerate(pages[: self.max_pages]):
                page_entities, page_text = self._predict_page(page, page_index)
                entities.extend(page_entities)
                if not document_text.strip() and page_text:
                    document_text += page_text + "\n"
            model_used = bool(entities)

        if not document_text.strip() and suffix in IMAGE_EXTENSIONS | {".pdf"}:
            pages = self._load_visual_pages(content, suffix)
            document_text = "\n".join(
                self._extract_image_text(page)
                for page in pages[: self.max_pages]
            )

        if not document_text.strip():
            raise UnsupportedDocumentError(
                "No readable text was found. Upload a clearer CV image or a text-based document."
            )

        return entities, document_text, model_used

    def _load_model(self) -> None:
        if self._model is not None:
            return
        with self._model_lock:
            if self._model is not None:
                return
            import torch
            from transformers import (
                AutoModelForTokenClassification,
                AutoProcessor,
                AutoTokenizer,
                LayoutLMv3ImageProcessor,
                LayoutLMv3Processor,
            )

            self._torch = torch
            try:
                self._processor = AutoProcessor.from_pretrained(
                    self.model_dir,
                    apply_ocr=False,
                    local_files_only=True,
                )
            except OSError:
                processor_config = self.model_dir / "processor_config.json"
                if not processor_config.is_file():
                    raise
                settings = json.loads(processor_config.read_text(encoding="utf-8"))
                image_settings = settings.get("image_processor")
                if not isinstance(image_settings, dict):
                    raise
                image_processor = LayoutLMv3ImageProcessor(**image_settings)
                tokenizer = AutoTokenizer.from_pretrained(
                    self.model_dir,
                    local_files_only=True,
                )
                self._processor = LayoutLMv3Processor(
                    image_processor=image_processor,
                    tokenizer=tokenizer,
                )
            self._model = AutoModelForTokenClassification.from_pretrained(
                self.model_dir,
                local_files_only=True,
            )
            self._model.eval()

    def _predict_page(self, image: Any, page_index: int = 0) -> tuple[list[EntitySpan], str]:
        self._load_model()
        words, pixel_boxes = self._ocr_words(image)
        if not words:
            return [], ""

        width, height = image.size
        boxes = [self._normalize_box(box, width, height) for box in pixel_boxes]
        encoding = self._processor(
            images=image,
            text=words,
            boxes=boxes,
            truncation=True,
            padding="max_length",
            max_length=512,
            return_tensors="pt",
        )
        word_ids = encoding.word_ids(batch_index=0)
        inputs = {
            key: value
            for key, value in encoding.items()
            if isinstance(value, self._torch.Tensor)
        }
        with self._torch.no_grad():
            probabilities = self._torch.softmax(
                self._model(**inputs).logits[0],
                dim=-1,
            )

        spans: list[EntitySpan] = []
        current_label: str | None = None
        current_words: list[str] = []
        current_scores: list[float] = []
        current_boxes: list[list[int]] = []
        current_order = 0
        seen_word_ids: set[int] = set()

        def flush() -> None:
            nonlocal current_label, current_words, current_scores
            nonlocal current_boxes, current_order
            if current_label and current_words:
                spans.append(
                    EntitySpan(
                        current_label,
                        " ".join(current_words),
                        sum(current_scores) / len(current_scores),
                        page_index,
                        self._union_box(current_boxes),
                        current_order,
                    )
                )
            current_label = None
            current_words = []
            current_scores = []
            current_boxes = []
            current_order = 0

        for token_index, word_id in enumerate(word_ids):
            if word_id is None or word_id in seen_word_ids or word_id >= len(words):
                continue
            seen_word_ids.add(word_id)
            token_probabilities = probabilities[token_index]
            prediction_id = int(token_probabilities.argmax())
            confidence = float(token_probabilities.max())
            label = self._model.config.id2label[prediction_id]

            if label == "O" or confidence < 0.45:
                flush()
                continue

            bio, entity = label.split("-", 1)
            if bio == "B" or current_label != entity:
                flush()
                current_label = entity
            if not current_boxes:
                # OCR word index is the reading order; use the span's first word.
                current_order = word_id
            current_words.append(words[word_id])
            current_scores.append(confidence)
            current_boxes.append(boxes[word_id])
        flush()
        return spans, " ".join(words)

    @staticmethod
    def _union_box(boxes: list[list[int]]) -> tuple[int, int, int, int] | None:
        if not boxes:
            return None
        return (
            min(box[0] for box in boxes),
            min(box[1] for box in boxes),
            max(box[2] for box in boxes),
            max(box[3] for box in boxes),
        )

    def _ocr_words(self, image: Any) -> tuple[list[str], list[list[float]]]:
        """Word boxes for the model path: PaddleOCR (train-consistent) or Tesseract."""
        if self.ocr_engine == "paddle":
            try:
                paddle_words = self._ocr_words_paddle(image)
                if paddle_words is not None and paddle_words[0]:
                    return paddle_words
            except Exception:
                LOGGER.warning("PaddleOCR failed; falling back to Tesseract.", exc_info=True)
        return self._ocr_words_tesseract(image)

    def _ocr_words_paddle(self, image: Any) -> tuple[list[str], list[list[float]]] | None:
        """OCR via PaddleOCR + the notebook's char-proportional word split."""
        engine = _load_paddle_engine()
        if engine is None:
            return None
        import numpy as np

        array = np.array(image.convert("RGB"))[:, :, ::-1]  # RGB -> BGR for Paddle
        result = engine.predict(array)
        if not result:
            return [], []
        page = result[0]
        texts = list(page.get("rec_texts", []) or [])
        polys = page.get("rec_polys")
        if polys is None or len(polys) == 0:
            polys = page.get("dt_polys", []) or []
        words: list[str] = []
        boxes: list[list[float]] = []
        for text, quad in zip(texts, polys):
            for word, box in _split_line_to_words(text, _quad_to_aabb(quad)):
                words.append(word)
                boxes.append(box)
        return words, boxes

    def _ocr_words_tesseract(self, image: Any) -> tuple[list[str], list[list[float]]]:
        completed = self._run_tesseract(image, "tsv")
        if completed.returncode != 0:
            raise UnsupportedDocumentError(
                "The CV image could not be read by OCR."
            )

        words: list[str] = []
        boxes: list[list[float]] = []
        rows = csv.DictReader(io.StringIO(completed.stdout), delimiter="\t")
        for row in rows:
            text = (row.get("text") or "").strip()
            if not text:
                continue
            try:
                left = float(row["left"])
                top = float(row["top"])
                width = float(row["width"])
                height = float(row["height"])
            except (KeyError, TypeError, ValueError):
                continue
            words.append(text)
            boxes.append([left, top, left + width, top + height])
        return words, boxes

    def _extract_text(self, content: bytes, suffix: str) -> str:
        if suffix == ".pdf":
            import fitz

            document = fitz.open(stream=content, filetype="pdf")
            return "\n".join(
                page.get_text("text")
                for page in list(document)[: self.max_pages]
            )
        if suffix == ".docx":
            from docx import Document

            document = Document(io.BytesIO(content))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            table_lines = [
                " | ".join(cell.text for cell in row.cells)
                for table in document.tables
                for row in table.rows
            ]
            return "\n".join(paragraphs + table_lines)
        if suffix == ".doc":
            with tempfile.NamedTemporaryFile(suffix=".doc") as temporary_file:
                temporary_file.write(content)
                temporary_file.flush()
                completed = subprocess.run(
                    ["antiword", temporary_file.name],
                    check=False,
                    capture_output=True,
                    text=True,
                    timeout=30,
                )
                if completed.returncode != 0:
                    raise UnsupportedDocumentError(
                        "The DOC file could not be read. Convert it to PDF or DOCX and try again."
                    )
                return completed.stdout
        return ""

    def _extract_image_text(self, image: Any) -> str:
        if not self.image_ocr_available:
            raise UnsupportedDocumentError(
                "Image OCR is unavailable in the AI service."
            )

        completed = self._run_tesseract(image, "text")
        if completed.returncode != 0:
            raise UnsupportedDocumentError(
                "The CV image could not be read by OCR."
            )
        return completed.stdout.strip()

    def _run_tesseract(self, image: Any, output_format: str) -> subprocess.CompletedProcess[str]:
        with tempfile.NamedTemporaryFile(suffix=".png") as temporary_file:
            image.convert("RGB").save(temporary_file.name, format="PNG")
            command = [
                "tesseract",
                temporary_file.name,
                "stdout",
                "-l",
                self.ocr_languages,
                "--psm",
                "6",
            ]
            if output_format == "tsv":
                command.append("tsv")
            return subprocess.run(
                command,
                check=False,
                capture_output=True,
                text=True,
                timeout=60,
            )

    def _load_visual_pages(self, content: bytes, suffix: str) -> list[Any]:
        if suffix in IMAGE_EXTENSIONS:
            return [self._load_image(content)]

        import fitz
        from PIL import Image

        document = fitz.open(stream=content, filetype="pdf")
        pages = []
        for page in list(document)[: self.max_pages]:
            pixmap = page.get_pixmap(matrix=fitz.Matrix(1.6, 1.6), alpha=False)
            pages.append(Image.open(io.BytesIO(pixmap.tobytes("png"))).convert("RGB"))
        return pages

    @staticmethod
    def _load_image(content: bytes) -> Any:
        from PIL import Image

        return Image.open(io.BytesIO(content)).convert("RGB")

    @staticmethod
    def _normalize_box(box: list[float], width: int, height: int) -> list[int]:
        x1, x2 = sorted([box[0], box[2]])
        y1, y2 = sorted([box[1], box[3]])

        def normalize(value: float, maximum: int) -> int:
            return min(1000, max(0, round(1000 * value / max(maximum, 1))))

        return [
            normalize(x1, width),
            normalize(y1, height),
            normalize(x2, width),
            normalize(y2, height),
        ]
